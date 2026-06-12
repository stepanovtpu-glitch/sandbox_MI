from datetime import datetime
from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.database import DB_DIR
from app.pilot_readiness import get_pilot_readiness

REPORTS_DIR = DB_DIR / 'reports'


def _report_name(suffix: Literal['pdf', 'docx']) -> Path:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return REPORTS_DIR / f'pilot_readiness_{stamp}.{suffix}'


def _status_label(status: str) -> str:
    return {
        'pilot_ready': 'Готов к ограниченному пилоту',
        'pilot_limited': 'Ограниченная готовность',
        'not_ready': 'Не готов к пилоту',
        'pass': 'Выполнено',
        'partial': 'Частично',
        'fail': 'Не выполнено',
    }.get(status, status)


def _checks_rows(readiness: dict[str, Any]) -> list[list[str]]:
    rows = [['Критерий', 'Статус', 'Вес', 'Балл', 'Подробности']]
    for check in readiness['checks']:
        rows.append([
            check['title'],
            _status_label(check['status']),
            str(check['weight']),
            str(check['score']),
            check['details'],
        ])
    return rows


def generate_readiness_pdf_report() -> Path:
    readiness = get_pilot_readiness()
    path = _report_name('pdf')
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=34, leftMargin=34, topMargin=34, bottomMargin=34)
    styles = getSampleStyleSheet()
    story = [
        Paragraph('GasMeter Pro — отчёт готовности к пилотной эксплуатации', styles['Title']),
        Paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}', styles['BodyText']),
        Spacer(1, 10),
        Paragraph(f'Итоговый статус: {_status_label(readiness["status"])}', styles['Heading2']),
        Paragraph(f'Готовность: {readiness["readiness_percent"]}% ({readiness["score"]} / {readiness["max_score"]})', styles['BodyText']),
        Paragraph(readiness['summary'], styles['BodyText']),
        Spacer(1, 12),
        Paragraph('Контрольный чек-лист', styles['Heading2']),
    ]
    table = Table(_checks_rows(readiness), colWidths=[150, 70, 35, 35, 205])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#DDE9F2')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.45, colors.HexColor('#6A7C86')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTSIZE', (0, 0), (-1, -1), 7),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph('Примечание', styles['Heading2']))
    story.append(Paragraph('Отчёт отражает техническую и организационную готовность пилотной версии. Перед промышленным применением обязательна ручная метрологическая сверка реальных МИ и утверждение матрицы МИ → формула → тест.', styles['BodyText']))
    doc.build(story)
    return path


def generate_readiness_docx_report() -> Path:
    readiness = get_pilot_readiness()
    path = _report_name('docx')
    document = Document()
    document.add_heading('GasMeter Pro — отчёт готовности к пилотной эксплуатации', level=1)
    document.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
    document.add_heading('Итог', level=2)
    document.add_paragraph(f'Статус: {_status_label(readiness["status"])}')
    document.add_paragraph(f'Готовность: {readiness["readiness_percent"]}% ({readiness["score"]} / {readiness["max_score"]})')
    document.add_paragraph(readiness['summary'])
    document.add_heading('Контрольный чек-лист', level=2)
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['Критерий', 'Статус', 'Вес', 'Балл', 'Подробности']
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in _checks_rows(readiness)[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.add_heading('Примечание', level=2)
    note = document.add_paragraph('Отчёт отражает техническую и организационную готовность пилотной версии. Перед промышленным применением обязательна ручная метрологическая сверка реальных МИ и утверждение матрицы МИ → формула → тест.')
    for run in note.runs:
        run.font.size = Pt(9)
    document.save(str(path))
    return path
