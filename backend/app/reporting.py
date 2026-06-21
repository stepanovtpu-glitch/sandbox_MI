from datetime import datetime
from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.database import DB_DIR
from app.method_library import list_method_versions
from app.pdf_fonts import apply_cyrillic_styles, register_pdf_fonts
from app.schemas import CalculationRequest, CalculationResult

REPORTS_DIR = DB_DIR / 'reports'


def _safe(value: str) -> str:
    return ''.join(ch if ch.isalnum() or ch in ('-', '_') else '_' for ch in value)[:80]


def _report_name(method_id: str | None, suffix: Literal['pdf', 'docx']) -> Path:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return REPORTS_DIR / f'calculation_protocol_{_safe(method_id or "method")}_{stamp}.{suffix}'


def _active_method_version(method_id: str | None) -> dict[str, Any] | None:
    if not method_id:
        return None
    versions = list_method_versions(method_id)
    return next((version for version in versions if version.get('status') == 'active'), versions[0] if versions else None)


def _document_summary(version: dict[str, Any] | None) -> tuple[str, str, str]:
    document = version.get('document') if version else None
    if not document:
        return '—', '—', '—'
    return (
        document.get('file_name') or '—',
        document.get('sha256') or '—',
        document.get('storage_path') or '—',
    )


def _rows_from_request(request: CalculationRequest, result: CalculationResult) -> list[list[str]]:
    method = request.method
    version = _active_method_version(method.mi_id if method else None)
    document_name, document_sha256, document_path = _document_summary(version)
    conclusion = (
        'Расчёт соответствует выбранной МИ.'
        if result.status == 'pass'
        else 'Расчёт не соответствует выбранной МИ. Требуется корректировка конфигурации или выбор другой МИ.'
        if result.status == 'fail'
        else 'Расчёт требует дополнительной проверки.'
    )
    rows = [
        ['Методика', method.title if method else 'не выбрана'],
        ['Регистрационный номер', method.registration_number if method else '—'],
        ['Версия МИ в библиотеке', version.get('version_id', '—') if version else '—'],
        ['Статус версии МИ', version.get('status', '—') if version else '—'],
        ['Скан-копия МИ', document_name],
        ['SHA-256 скан-копии', document_sha256],
        ['Путь хранения скан-копии', document_path],
        ['Контрольных примеров МИ', str(len(version.get('test_cases', []))) if version else '0'],
        ['Диапазон Q, м³/ч', f'{request.line.q_min}–{request.line.q_max}'],
        ['Диапазон P, МПа', f'{request.line.p_min_mpa}–{request.line.p_max_mpa}'],
        ['Диапазон T, °C', f'{request.line.t_min_c}–{request.line.t_max_c}'],
        ['Расчётный шаблон', request.calculation_template.value if request.calculation_template else 'MANUAL_QUADRATURE'],
        ['Итоговая U / δΣ, %', str(result.delta_total)],
        ['Предел МИ, %', str(result.limit) if result.limit is not None else '—'],
        ['Статус', result.status],
        ['Заключение', conclusion],
    ]
    if request.instruments:
        rows.append(['Состав СИ', ''])
        for instrument in request.instruments:
            rows.append([
                f'СИ: {instrument.type.value}',
                (
                    f'{instrument.name}; зав. № {instrument.serial_number or "—"}; '
                    f'δ={instrument.error_percent if instrument.error_percent is not None else "—"}%; '
                    f'диапазон={instrument.range_min if instrument.range_min is not None else "—"}–{instrument.range_max if instrument.range_max is not None else "—"} {instrument.range_unit or ""}; '
                    f'свидетельство={instrument.certificate_number or "—"}; поверка до={instrument.calibration_due or "—"}'
                ),
            ])
    return rows


def _apply_table_style(table: Table, first_col_bg: str = '#EAF7F5') -> None:
    regular, _ = register_pdf_fonts()
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor(first_col_bg)),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#6A7C86')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), regular),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
    ]))


def _pdf_rows(rows: list[list[str]], styles) -> list[list[Paragraph]]:
    return [[Paragraph(str(cell), styles['BodyText']) for cell in row] for row in rows]


def generate_pdf_report(request: CalculationRequest, result: CalculationResult) -> Path:
    path = _report_name(request.method.mi_id if request.method else None, 'pdf')
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    apply_cyrillic_styles(styles)
    story = [
        Paragraph('GasMeter Pro — протокол расчёта', styles['Title']),
        Paragraph('Автоматизированный расчёт погрешности / расширенной неопределённости измерений', styles['BodyText']),
        Paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}', styles['BodyText']),
        Spacer(1, 12),
    ]
    table = Table(_pdf_rows(_rows_from_request(request, result), styles), colWidths=[155, 340])
    _apply_table_style(table)
    story.append(table)
    story.append(Spacer(1, 14))
    story.append(Paragraph('Вклады составляющих', styles['Heading2']))
    contrib_rows = [['Код', 'Наименование', 'Значение', 'Взвешенное', 'Доля, %']]
    for item in result.contributions:
        contrib_rows.append([item.code, item.label, str(item.value), str(item.weighted_value), str(item.share_percent)])
    contrib_table = Table(_pdf_rows(contrib_rows, styles), colWidths=[60, 235, 65, 75, 60])
    regular, _ = register_pdf_fonts()
    contrib_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#DDE9F2')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#6A7C86')),
        ('FONTNAME', (0, 0), (-1, -1), regular),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.append(contrib_table)
    story.append(Spacer(1, 14))
    story.append(Paragraph('Аудит расчёта', styles['Heading2']))
    for row in result.audit_log:
        story.append(Paragraph(row, styles['Code']))
    doc.build(story)
    return path


def generate_docx_report(request: CalculationRequest, result: CalculationResult) -> Path:
    path = _report_name(request.method.mi_id if request.method else None, 'docx')
    document = Document()
    document.add_heading('GasMeter Pro — протокол расчёта', level=1)
    document.add_paragraph('Автоматизированный расчёт погрешности / расширенной неопределённости измерений')
    document.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for key, value in _rows_from_request(request, result):
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    document.add_heading('Вклады составляющих', level=2)
    ctable = document.add_table(rows=1, cols=5)
    ctable.style = 'Table Grid'
    headers = ['Код', 'Наименование', 'Значение', 'Взвешенное', 'Доля, %']
    for idx, header in enumerate(headers):
        ctable.rows[0].cells[idx].text = header
    for item in result.contributions:
        cells = ctable.add_row().cells
        cells[0].text = item.code
        cells[1].text = item.label
        cells[2].text = str(item.value)
        cells[3].text = str(item.weighted_value)
        cells[4].text = str(item.share_percent)
    document.add_heading('Аудит расчёта', level=2)
    for row in result.audit_log:
        p = document.add_paragraph(row)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
    document.save(str(path))
    return path
