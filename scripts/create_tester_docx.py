from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIRS = [
    ROOT / 'docs' / 'tester_package',
    ROOT / 'release' / 'GasMeterPro' / 'docs',
]

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
HEADER_FILL = 'E8EEF5'
LIGHT_FILL = 'F4F6F9'
BORDER = 'B7C3D0'


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths_dxa)))
    tbl_w.set(qn('w:type'), 'dxa')

    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), BORDER)


def add_footer(section, label: str) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(label)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(89, 89, 89)


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section, 'GasMeter Pro alpha package')

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 14, 7),
        ('Heading 3', 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(3)
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(89, 89, 89)
    return doc


def add_note(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f'{label}: ')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(text)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    set_borders(table)
    for idx, (key, value) in enumerate(rows):
        if idx > 0:
            table.add_row()
        row = table.rows[idx]
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_shading(row.cells[0], HEADER_FILL)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            row.cells[0].paragraphs[0].runs[0].bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_borders(table)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        cell.paragraphs[0].runs[0].bold = True
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_description_doc() -> Document:
    doc = setup_doc(
        'GasMeter Pro. Описание программы и инструкция тестировщика',
        'Версия для альфа-тестирования. Portable-сборка для Windows без установки Python, Node.js, Tesseract и Poppler.',
    )
    add_note(
        doc,
        'Назначение',
        'GasMeter Pro предназначена для подбора и проверки методик измерений газа, ведения библиотеки МИ и СИ, расчета погрешности / расширенной неопределенности и формирования протоколов.',
    )

    doc.add_heading('1. Что входит в комплект', level=1)
    add_kv_table(doc, [
        ('GasMeterPro.exe', 'Локальный сервер приложения. Запускает backend, раздает интерфейс и API.'),
        ('start-gasmeter.bat', 'Основной файл запуска для тестировщика. Открывает браузер и запускает программу.'),
        ('data', 'Рабочая папка пользователя: база SQLite, загруженные PDF МИ, OCR-тексты, отчеты и история расчетов.'),
        ('_internal', 'Встроенные зависимости: Python runtime, frontend, Tesseract OCR, модели rus/eng, Poppler.'),
        ('docs', 'Документы для тестировщика: описание и программа тестирования.'),
    ])

    doc.add_heading('2. Как запустить', level=1)
    add_numbered(doc, [
        'Распаковать архив GasMeterPro-portable-alpha.zip в отдельную папку без прав администратора.',
        'Открыть файл start-gasmeter.bat двойным щелчком.',
        'Дождаться открытия браузера по адресу http://127.0.0.1:8000/.',
        'Если браузер не открылся автоматически, открыть этот адрес вручную.',
        'Не закрывать окно запуска, пока идет тестирование: при закрытии окна локальный сервер остановится.',
    ])
    add_note(doc, 'Важно', 'Программа работает локально на компьютере тестировщика. Данные не отправляются в интернет, если пользователь сам не передает файлы разработчику.')

    doc.add_heading('3. Основные разделы программы', level=1)
    add_table(
        doc,
        ['Раздел', 'Назначение', 'Что проверить'],
        [
            ['Конструктор УУГ', 'Ввод параметров измерительной линии, СИ и расчет неопределенности.', 'Пересчет после изменения Q/P/T и ошибок СИ, статус pass/warn/fail.'],
            ['Подбор МИ', 'Подбор подходящей методики по технологическому режиму.', 'Рекомендации, причины выбора, применение выбранной МИ.'],
            ['База СИ', 'Просмотр средств измерений и применение приборов в расчет.', 'Фильтрация по типам, статусы поверки, применение прибора.'],
            ['Добавление МИ', 'Полный цикл ввода новой МИ: PDF, OCR, правка, расчетная проверка, подтверждение.', 'Главный раздел для тестирования новых методик.'],
            ['Библиотека МИ', 'Версии методик, документы, контрольные примеры, история правок.', 'Загрузка PDF, OCR, подтверждение, список версий.'],
            ['Отчеты', 'Формирование протоколов и отчета готовности.', 'PDF/DOCX отчеты, читаемость кириллицы, наличие SHA и исходных данных.'],
            ['Настройки', 'Состояние backend и сведения о конфигурации.', 'Health/status, путь к базе, количество МИ и СИ.'],
        ],
        [1700, 3800, 3860],
    )

    doc.add_heading('4. Рабочий сценарий тестировщика', level=1)
    add_numbered(doc, [
        'Открыть программу через start-gasmeter.bat.',
        'В разделе “Конструктор УУГ” изменить параметры линии и убедиться, что расчет обновляется.',
        'В разделе “Подбор МИ” подобрать методику для заданного режима и применить ее.',
        'В разделе “Добавление МИ” выбрать существующую или создать новую МИ.',
        'Загрузить PDF методики, выполнить SHA-проверку и OCR.',
        'Сравнить OCR с карточкой МИ, исправить ошибочно распознанные поля, сохранить версию.',
        'Выполнить расчетную проверку и добавить контрольный пример.',
        'Подтвердить проверку МИ и убедиться, что статус отображается в чеклисте готовности.',
        'Сформировать PDF/DOCX протокол и проверить читаемость текста.',
    ])

    doc.add_heading('5. Что считать ошибкой', level=1)
    add_bullets(doc, [
        'Программа не запускается или браузер не открывает локальный адрес.',
        'После ввода чисел расчет не обновляется или показывает очевидно неверный статус.',
        'PDF не загружается, SHA-проверка падает без понятного сообщения.',
        'OCR зависает, не показывает результат или ломает ранее заполненную карточку МИ.',
        'В протоколах есть нечитаемая кириллица, пустые таблицы, обрезанный текст или неправильные значения.',
        'После перезапуска исчезают добавленные МИ, PDF, история расчетов или подтверждения.',
    ])

    doc.add_heading('6. Как передавать результат тестирования', level=1)
    add_kv_table(doc, [
        ('Минимум', 'Скриншот ошибки, краткое описание действия, при котором она возникла, и дата/время.'),
        ('Желательно', 'Архив папки data после теста, если ошибка связана с базой, PDF, OCR или историей расчетов.'),
        ('Не передавать', 'Конфиденциальные реальные документы без разрешения владельца данных.'),
    ])
    return doc


def build_test_program_doc() -> Document:
    doc = setup_doc(
        'GasMeter Pro. Программа альфа-тестирования',
        'Рабочий план проверки portable-сборки, основных функций, OCR, базы МИ/СИ, расчетов и отчетов.',
    )
    add_note(
        doc,
        'Цель тестирования',
        'Проверить, что программа запускается на ПК без установленного Python/Node.js, сохраняет данные локально, корректно ведет МИ/СИ и позволяет выполнить расчетный цикл от PDF методики до протокола.',
    )

    doc.add_heading('1. Область тестирования', level=1)
    add_table(
        doc,
        ['Блок', 'Проверяемый результат', 'Критерий приемки'],
        [
            ['Запуск', 'Portable-сборка стартует через start-gasmeter.bat.', 'Браузер открывает http://127.0.0.1:8000/, health/status доступны.'],
            ['Расчет', 'Расчет U/δΣ обновляется при изменении параметров.', 'Нет ошибок API, итоговый статус соответствует введенным пределам.'],
            ['Подбор МИ', 'МИ подбирается по Q/P/T и шаблону.', 'Есть ранжированный список и понятные причины выбора.'],
            ['Добавление МИ', 'Технолог проходит путь PDF → OCR → правка → проверка → подтверждение.', 'Чеклист готовности показывает выполненные этапы.'],
            ['Библиотека МИ', 'Версии, документы и контрольные примеры сохраняются.', 'После перезапуска данные остаются в базе.'],
            ['Отчеты', 'PDF/DOCX формируются и читаются.', 'Кириллица читаема, таблицы не обрезаны, SHA и параметры присутствуют.'],
            ['Надежность', 'Некорректные данные не ломают приложение.', 'Показывается понятное сообщение, интерфейс продолжает работать.'],
        ],
        [1800, 3900, 3660],
    )

    doc.add_heading('2. Подготовка теста', level=1)
    add_numbered(doc, [
        'Создать отдельную папку для теста, например C:\\GasMeterProAlpha.',
        'Распаковать GasMeterPro-portable-alpha.zip.',
        'Запустить start-gasmeter.bat.',
        'Подготовить 2-3 PDF методики: один хороший скан, один многостраничный PDF, один проблемный или низкокачественный скан.',
        'Подготовить тестовые параметры линии: Q, P, T, DN, погрешности СИ и ожидаемый результат контрольного примера, если он известен.',
    ])

    doc.add_heading('3. Сценарии тестирования', level=1)
    add_table(
        doc,
        ['ID', 'Сценарий', 'Действия', 'Ожидаемый результат', 'Факт / замечание'],
        [
            ['T-01', 'Первый запуск', 'Запустить start-gasmeter.bat.', 'Открылась программа, ошибок в интерфейсе нет.', ''],
            ['T-02', 'Перезапуск', 'Закрыть окно запуска, открыть снова.', 'Данные и список МИ/СИ сохранились.', ''],
            ['T-03', 'Расчет линии', 'Изменить Q/P/T и ошибки СИ в конструкторе.', 'Итог U/δΣ пересчитался, статус обновился.', ''],
            ['T-04', 'Подбор МИ', 'Ввести технологический режим и применить рекомендованную МИ.', 'МИ выбрана в конструкторе, параметры применены.', ''],
            ['T-05', 'Загрузка PDF МИ', 'В “Добавление МИ” загрузить PDF.', 'PDF отображается в версии, SHA рассчитан.', ''],
            ['T-06', 'OCR методики', 'Запустить OCR по PDF.', 'Появился статус OCR и извлеченные поля.', ''],
            ['T-07', 'Правка распознавания', 'Исправить неверные Q/P/T/δ в карточке и сохранить версию.', 'Исправленные значения остались после обновления страницы.', ''],
            ['T-08', 'Контроль расчета МИ', 'Запустить расчетную проверку и добавить контрольный пример.', 'Результат pass либо понятная причина fail.', ''],
            ['T-09', 'Подтверждение МИ', 'Нажать подтверждение проверки МИ.', 'В чеклисте готовности ручная сверка отмечена выполненной.', ''],
            ['T-10', 'Протокол PDF/DOCX', 'Сформировать протокол расчета.', 'Файл скачался, кириллица читаема, таблицы не обрезаны.', ''],
            ['T-11', 'История', 'Сохранить расчет, открыть журнал.', 'Расчет виден в истории и загружается обратно.', ''],
            ['T-12', 'Негативный тест', 'Ввести Q max меньше Q min или загрузить неподходящий файл.', 'Приложение не падает, показывает ошибку.', ''],
        ],
        [700, 1700, 2850, 2700, 1410],
    )

    doc.add_heading('4. Матрица приемки альфа-версии', level=1)
    add_table(
        doc,
        ['Проверка', 'Минимальный критерий', 'Статус'],
        [
            ['Запуск без Python/Node.js', 'Программа запускается на чистом ПК через BAT-файл.', '□ pass / □ fail'],
            ['Локальное хранение', 'Папка data содержит базу и не теряется при перезапуске.', '□ pass / □ fail'],
            ['Основной расчет', 'Конструктор выполняет расчет без API-ошибок.', '□ pass / □ fail'],
            ['Добавление МИ', 'Новая или существующая МИ проходит загрузку PDF и сохранение версии.', '□ pass / □ fail'],
            ['OCR', 'OCR запускается и возвращает хотя бы текст/статус качества.', '□ pass / □ fail'],
            ['Ручная правка', 'Ошибочные OCR-значения можно исправить и сохранить.', '□ pass / □ fail'],
            ['Контрольный пример', 'Пример добавляется и запускается.', '□ pass / □ fail'],
            ['Отчеты', 'PDF/DOCX формируются без нечитаемой кириллицы.', '□ pass / □ fail'],
        ],
        [3000, 4700, 1660],
    )

    doc.add_heading('5. Форма отчета тестировщика', level=1)
    add_table(
        doc,
        ['Поле', 'Заполняет тестировщик'],
        [
            ['ФИО / подразделение', ''],
            ['Дата тестирования', ''],
            ['Windows версия', ''],
            ['Путь установки программы', ''],
            ['PDF методики, использованные в тесте', ''],
            ['Что прошло успешно', ''],
            ['Найденные ошибки', ''],
            ['Предложения по улучшению', ''],
            ['Общий вывод', '□ можно передавать дальше / □ нужны исправления / □ блокирующие ошибки'],
        ],
        [3000, 6360],
    )

    doc.add_heading('6. Приоритеты дефектов', level=1)
    add_kv_table(doc, [
        ('P1 блокирующий', 'Программа не запускается, невозможно открыть интерфейс, потеря данных после перезапуска, критически неверный расчет.'),
        ('P2 высокий', 'Нельзя загрузить PDF, OCR не работает на большинстве файлов, невозможно подтвердить МИ или сформировать протокол.'),
        ('P3 средний', 'Ошибки интерфейса, неудобная последовательность действий, отдельные поля не сохраняются.'),
        ('P4 низкий', 'Опечатки, косметика, пожелания по удобству и формулировкам.'),
    ])
    return doc


def save_doc(doc: Document, filename: str) -> None:
    for out_dir in OUT_DIRS:
        out_dir.mkdir(parents=True, exist_ok=True)
        doc.save(out_dir / filename)


def main() -> None:
    save_doc(build_description_doc(), 'GasMeter_Pro_Описание_и_инструкция_тестировщика.docx')
    save_doc(build_test_program_doc(), 'GasMeter_Pro_Программа_альфа_тестирования.docx')


if __name__ == '__main__':
    main()
